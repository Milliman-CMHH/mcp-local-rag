import hashlib
from dataclasses import dataclass
from pathlib import Path

import aiofiles
import pymupdf  # type: ignore[import-untyped]
import pymupdf4llm  # type: ignore[import-untyped]
from docx import Document

from mcp_local_rag.config import SUPPORTED_EXTENSIONS
from mcp_local_rag.processing.ocr import needs_ocr, ocr_pdf_pages


@dataclass
class ExtractedDocument:
    file_path: str
    file_hash: str
    content: str
    file_type: str
    page_count: int | None = None


def compute_file_hash(file_path: Path) -> str:
    with open(file_path, "rb") as f:
        return hashlib.file_digest(f, "sha256").hexdigest()


async def extract_pdf(file_path: Path) -> ExtractedDocument:
    file_hash = compute_file_hash(file_path)

    doc = pymupdf.open(str(file_path))
    page_count: int = len(doc)

    # First try standard text extraction
    md_text: str = pymupdf4llm.to_markdown(str(file_path))  # type: ignore[assignment]

    # If extracted text is insufficient, use OCR
    if needs_ocr(md_text, page_count):
        md_text = ocr_pdf_pages(doc, dpi=300)

    doc.close()

    return ExtractedDocument(
        file_path=str(file_path.resolve()),
        file_hash=file_hash,
        content=md_text,
        file_type="pdf",
        page_count=page_count,
    )


async def extract_docx(file_path: Path) -> ExtractedDocument:
    file_hash = compute_file_hash(file_path)
    doc = Document(str(file_path))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return ExtractedDocument(
        file_path=str(file_path.resolve()),
        file_hash=file_hash,
        content="\n\n".join(paragraphs),
        file_type="docx",
    )


async def extract_plaintext(file_path: Path) -> ExtractedDocument:
    file_hash = compute_file_hash(file_path)

    async with aiofiles.open(file_path, "r", encoding="utf-8", errors="replace") as f:
        content = await f.read()

    return ExtractedDocument(
        file_path=str(file_path.resolve()),
        file_hash=file_hash,
        content=content,
        file_type="plaintext",
    )


async def extract_document(file_path: Path) -> ExtractedDocument:
    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")

    file_type = SUPPORTED_EXTENSIONS[suffix]

    if file_type == "pdf":
        return await extract_pdf(file_path)
    elif file_type == "docx":
        return await extract_docx(file_path)
    else:
        return await extract_plaintext(file_path)


def is_supported_file(file_path: Path) -> bool:
    return file_path.suffix.lower() in SUPPORTED_EXTENSIONS
